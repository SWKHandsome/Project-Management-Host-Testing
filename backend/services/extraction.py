"""
Data extraction service
Extracts student information from document content
"""
import re
from PyPDF2 import PdfReader
from docx import Document
import io

from config import Config

class DataExtractor:
    """Extract student data from various file formats"""
    
    def extract_from_content(self, file_content, filename):
        """Extract student info from file content"""
        try:
            file_ext = filename.lower().split('.')[-1]
            
            text = ""
            if file_ext == 'pdf':
                text = self.extract_from_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                text = self.extract_from_docx(file_content)
            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
            
            # Extract student information from text
            return self.parse_student_info(text)
            
        except Exception as e:
            print(f"  ⚠ Error extracting from content: {e}")
            return {'student_id': None, 'student_name': None}
    
    def extract_from_pdf(self, pdf_content):
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(pdf_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            # Read first 3 pages (usually contains student info)
            for page_num in range(min(3, len(reader.pages))):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            print(f"  ⚠ Error reading PDF: {e}")
            return ""
    
    def extract_from_docx(self, docx_content):
        """Extract text from DOCX"""
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            text = ""
            # Read first 10 paragraphs (usually contains student info)
            for para in doc.paragraphs[:10]:
                text += para.text + "\n"
            
            return text
            
        except Exception as e:
            print(f"  ⚠ Error reading DOCX: {e}")
            return ""
    
    def parse_student_info(self, text):
        """Parse student ID and name from text"""
        student_id = None
        student_name = None
        
        # Extract student ID using patterns
        for pattern in Config.STUDENT_ID_PATTERNS:
            match = re.search(pattern, text)
            if match:
                student_id = match.group(0)
                break
        
        # Extract student name using patterns
        for pattern in Config.NAME_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                student_name = match.group(1).strip()
                break
        
        # If no pattern match, try to find name after common keywords
        if not student_name:
            keywords = ['name', 'student', 'by', 'author', 'prepared by']
            for keyword in keywords:
                pattern = rf'{keyword}\s*[:\-]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    student_name = match.group(1).strip()
                    break
        
        return {
            'student_id': student_id,
            'student_name': student_name
        }
    
    def extract_assignment_content(self, file_content, filename):
        """Extract full assignment content for evaluation"""
        try:
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                return self.extract_full_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                return self.extract_full_docx(file_content)
            elif file_ext == 'txt':
                return file_content.decode('utf-8', errors='ignore')
            
            return ""
            
        except Exception as e:
            print(f"  ⚠ Error extracting assignment content: {e}")
            return ""
    
    def extract_full_pdf(self, pdf_content):
        """Extract all text from PDF"""
        try:
            pdf_file = io.BytesIO(pdf_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            return ""
    
    def extract_full_docx(self, docx_content):
        """Extract all text from DOCX"""
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            return text
            
        except Exception as e:
            return ""
